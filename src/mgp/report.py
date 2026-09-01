"""
Generazione dei report settimanali per il relatore in formato Word (.docx).

Perche' Markdown come sorgente e Word come output
-------------------------------------------------
Il testo del report si scrive in Markdown (`docs/report/AAAA-MM-GG_report.md`): e' testo
semplice, quindi si versiona, si confronta fra una settimana e l'altra e si riusa nei
capitoli della tesi. Il relatore riceve pero' un `.docx`, che e' il formato con cui e'
comodo commentare e annotare. Questo modulo converte il primo nel secondo.

E' un convertitore volutamente minimale: supporta il sottoinsieme di Markdown che serve a
un report metodologico, senza dipendere da pandoc o da altri strumenti esterni.

Sintassi supportata
-------------------
* titoli `#`, `##`, `###`, `####`;
* paragrafi, righe vuote;
* elenchi puntati (`- `, `* `) e numerati (`1. `), con un livello di annidamento;
* tabelle Markdown (`| a | b |` con riga di separazione);
* blocchi di codice delimitati da ``` (resi in Consolas, sfondo grigio);
* citazioni (`> `);
* linee orizzontali (`---`);
* formattazione inline: `**grassetto**`, `*corsivo*`, `` `codice` ``,
  link `[testo](url)` (resi come "testo (url)");
* immagini `![didascalia](file.png)` su riga propria (vedi `_immagine`).
"""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Emu, Pt, RGBColor

from mgp import config

# --------------------------------------------------------------------------------------
# Formattazione inline
# --------------------------------------------------------------------------------------
#: Un unico pattern che cattura, nell'ordine, codice inline, grassetto, corsivo e link.
#: L'ordine conta: `**testo**` va riconosciuto prima di `*testo*`.
_INLINE = re.compile(
    r"(`[^`]+`)"                 # codice
    r"|(\*\*[^*]+\*\*)"          # grassetto
    r"|(\*[^*]+\*)"              # corsivo
    r"|(\[[^\]]+\]\([^)]+\))"    # link
)


#: Sequenze di escape Markdown: una barra rovesciata davanti a un carattere di
#: punteggiatura significa "questo carattere e' letterale, non un delimitatore".
#: Nei report ricorre soprattutto in `K\*`, dove serve a impedire che l'asterisco
#: apra un corsivo.
_ESCAPE = re.compile(r"\\([\\`*_{}\[\]()#+\-.!])")

#: Base dell'area a uso privato di Unicode, dove il carattere protetto viene
#: temporaneamente spostato.
#:
#: Non basta togliere la barra rovesciata e lasciare il carattere dov'e': l'asterisco
#: di `K\*` resterebbe un delimitatore e due occorrenze sulla stessa riga verrebbero
#: lette come un corsivo. Il carattere va quindi **nascosto per intero** finche' la
#: formattazione inline non e' stata riconosciuta. La punteggiatura che Markdown
#: ammette dopo la barra e' tutta ASCII, quindi `_PUA + ord(c)` resta dentro l'area
#: a uso privato (E000-F8FF) e la trasformazione e' invertibile.
_PUA = 0xE000


def _proteggi_escape(testo: str) -> str:
    """Sposta i caratteri protetti da `\\` nell'area a uso privato di Unicode."""
    return _ESCAPE.sub(lambda m: chr(_PUA + ord(m.group(1))), testo)


def _scioglie_escape(pezzo: str, dentro_codice: bool = False) -> str:
    """
    Riporta i caratteri protetti al loro valore, dopo il riconoscimento inline.

    Dentro il codice inline si ripristina anche la barra rovesciata, perche' in
    Markdown gli escape non agiscono in un blocco di codice: chi ha scritto `\\*` fra
    apici inversi voleva vedere proprio quei due caratteri.
    """
    prefisso = "\\" if dentro_codice else ""
    return "".join(
        prefisso + chr(ord(c) - _PUA) if _PUA <= ord(c) < _PUA + 128 else c for c in pezzo
    )


def _aggiungi_testo(paragrafo, testo: str) -> None:
    """
    Aggiunge a un paragrafo Word il testo Markdown, traducendo la formattazione inline.

    Il testo viene spezzato in "run" (frammenti con formattazione omogenea), perche' in
    Word grassetto e corsivo sono proprieta' del run, non del paragrafo.
    """
    for pezzo in _INLINE.split(_proteggi_escape(testo)):
        if not pezzo:
            continue
        if pezzo.startswith("`") and pezzo.endswith("`"):
            run = paragrafo.add_run(_scioglie_escape(pezzo[1:-1], dentro_codice=True))
            run.font.name = "Consolas"
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(0xB0, 0x30, 0x30)
        elif pezzo.startswith("**") and pezzo.endswith("**"):
            paragrafo.add_run(_scioglie_escape(pezzo[2:-2])).bold = True
        elif pezzo.startswith("*") and pezzo.endswith("*"):
            paragrafo.add_run(_scioglie_escape(pezzo[1:-1])).italic = True
        elif pezzo.startswith("[") and "](" in pezzo:
            etichetta, url = pezzo[1:-1].split("](", 1)
            paragrafo.add_run(_scioglie_escape(etichetta))
            run = paragrafo.add_run(f" ({_scioglie_escape(url)})")
            run.font.size = Pt(8.5)
            run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
        else:
            paragrafo.add_run(_scioglie_escape(pezzo))


# --------------------------------------------------------------------------------------
# Immagini
# --------------------------------------------------------------------------------------
#: Immagine Markdown su riga propria: `![didascalia](percorso)`.
_IMMAGINE = re.compile(r"^!\[([^\]]*)\]\(([^)]+)\)$")


def _risolvi_figura(riferimento: str, sorgente: Path) -> Path | None:
    """
    Trova il file di una figura citata in un report.

    Si prova, nell'ordine: il percorso relativo alla cartella del report, quello relativo
    alla radice del progetto, e infine il nome nudo dentro `output/figure/`. L'ultima
    forma e' quella comoda da scrivere — `![...](15_errore_orario.png)` — perche' evita
    i `../..` che renderebbero il sorgente Markdown illeggibile.

    Word non sa incorporare un PDF: se il riferimento e' a un `.pdf` si cerca il `.png`
    di pari nome, che gli script di questo progetto producono sempre insieme.

    Returns
    -------
    Path | None
        Il file trovato, oppure None se non esiste: la conversione non si interrompe,
        perche' un report a meta' e' peggio di un report con una figura mancante
        segnalata (vedi `_immagine`).
    """
    candidati = [Path(riferimento)] if not Path(riferimento).is_absolute() else []
    if Path(riferimento).suffix.lower() == ".pdf":
        candidati.append(Path(riferimento).with_suffix(".png"))

    for candidato in list(candidati) or [Path(riferimento)]:
        for base in (sorgente.parent, config.PROJECT_ROOT, config.FIGURE_DIR):
            percorso = (base / candidato).resolve()
            if percorso.is_file():
                return percorso
    return None


def _immagine(doc: Document, didascalia: str, riferimento: str, sorgente: Path) -> None:
    """
    Inserisce una figura, ridimensionata alla larghezza della colonna di testo.

    L'immagine non viene mai ingrandita: una figura piu' stretta del testo resta della
    sua dimensione naturale, perche' scalarla in su la sfoca. La didascalia, se c'e',
    va sotto in corsivo piccolo.

    Se il file non si trova, al suo posto compare un segnaposto rosso con il nome
    cercato e viene stampato un avviso. Fallire in silenzio sarebbe il caso peggiore:
    il documento va al relatore, e una figura assente deve vedersi sia in fase di
    generazione sia nel documento stesso.
    """
    percorso = _risolvi_figura(riferimento, sorgente)
    if percorso is None:
        print(f"  ATTENZIONE: figura non trovata, '{riferimento}'", file=sys.stderr)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"[FIGURA MANCANTE: {riferimento}]")
        run.bold = True
        run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
        return

    sezione = doc.sections[0]
    larghezza_utile = sezione.page_width - sezione.left_margin - sezione.right_margin

    figura = doc.add_picture(str(percorso))
    if figura.width > larghezza_utile:
        # Si scala anche l'altezza a mano: fissare la sola larghezza deformerebbe.
        figura.height = Emu(int(figura.height * larghezza_utile / figura.width))
        figura.width = Emu(int(larghezza_utile))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    if didascalia:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Passa dal riconoscimento inline come ogni altro testo: una didascalia puo'
        # contenere un `K\*` da sciogliere o un nome di file fra apici inversi.
        _aggiungi_testo(p, didascalia)
        for run in p.runs:
            run.italic = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x50, 0x50, 0x50)


def _sfondo(paragrafo, colore_hex: str) -> None:
    """Applica un colore di sfondo a un paragrafo (usato per i blocchi di codice)."""
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), colore_hex)
    paragrafo._p.get_or_add_pPr().append(shd)


# --------------------------------------------------------------------------------------
# Conversione del documento
# --------------------------------------------------------------------------------------
def _blocco_codice(doc: Document, righe: list[str]) -> None:
    """Inserisce un blocco di codice come paragrafo monospaziato con sfondo grigio."""
    for riga in righe:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.left_indent = Pt(14)
        run = p.add_run(riga if riga.strip() else " ")
        run.font.name = "Consolas"
        run.font.size = Pt(9)
        _sfondo(p, "F2F2F2")


def _tabella(doc: Document, righe: list[str]) -> None:
    """
    Inserisce una tabella Markdown come tabella Word.

    Si assume che la prima riga sia l'intestazione e la seconda la riga di separazione
    (`|---|---|`), come da sintassi Markdown standard.
    """
    def celle(riga: str) -> list[str]:
        return [c.strip() for c in riga.strip().strip("|").split("|")]

    intestazione = celle(righe[0])
    corpo = [celle(r) for r in righe[2:] if r.strip()]

    tabella = doc.add_table(rows=1, cols=len(intestazione))
    tabella.style = "Light Grid Accent 1"
    for cella, testo in zip(tabella.rows[0].cells, intestazione):
        # La cella nasce con un paragrafo vuoto: ci si scrive dentro direttamente.
        run = cella.paragraphs[0].add_run(
            _scioglie_escape(_proteggi_escape(testo).replace("**", ""))
        )
        run.bold = True
        run.font.size = Pt(9.5)
    for riga in corpo:
        celle_word = tabella.add_row().cells
        for cella, testo in zip(celle_word, riga):
            p = cella.paragraphs[0]
            _aggiungi_testo(p, testo)
            for run in p.runs:
                run.font.size = Pt(9.5)
    doc.add_paragraph()


def markdown_to_docx(sorgente: Path, destinazione: Path) -> Path:
    """
    Converte un report Markdown in un documento Word.

    Parameters
    ----------
    sorgente : Path
        File `.md` del report (vedi `docs/report/`).
    destinazione : Path
        File `.docx` da produrre. Le cartelle mancanti vengono create.

    Returns
    -------
    Path
        Il path del file prodotto.

    Note
    ----
    Il documento usa gli stili predefiniti di Word (Title, Heading 1-3, List Bullet,
    List Number), cosi' il relatore puo' cambiarne l'aspetto senza toccare il contenuto e
    Word genera automaticamente la struttura di navigazione.
    """
    testo = sorgente.read_text(encoding="utf-8")
    doc = Document()

    # Corpo del testo leggibile a schermo e in stampa.
    stile = doc.styles["Normal"]
    stile.font.name = "Calibri"
    stile.font.size = Pt(11)

    righe = testo.splitlines()
    i = 0
    while i < len(righe):
        riga = righe[i]
        spoglia = riga.strip()

        # --- blocco di codice ---------------------------------------------------------
        if spoglia.startswith("```"):
            i += 1
            blocco: list[str] = []
            while i < len(righe) and not righe[i].strip().startswith("```"):
                blocco.append(righe[i])
                i += 1
            i += 1
            _blocco_codice(doc, blocco)
            doc.add_paragraph()
            continue

        # --- tabella ------------------------------------------------------------------
        if spoglia.startswith("|") and i + 1 < len(righe) and set(righe[i + 1].strip()) <= set("|-: "):
            blocco = []
            while i < len(righe) and righe[i].strip().startswith("|"):
                blocco.append(righe[i])
                i += 1
            _tabella(doc, blocco)
            continue

        # --- immagine -----------------------------------------------------------------
        # Va riconosciuta prima dei paragrafi: `![...](...)` non e' testo corrente.
        figura = _IMMAGINE.match(spoglia)
        if figura:
            _immagine(doc, figura.group(1).strip(), figura.group(2).strip(), sorgente)
            i += 1
            continue

        # --- riga vuota / separatore --------------------------------------------------
        if not spoglia:
            i += 1
            continue
        if set(spoglia) <= set("-") and len(spoglia) >= 3:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            _sfondo(p, "DDDDDD")
            i += 1
            continue

        # --- titoli -------------------------------------------------------------------
        if spoglia.startswith("#"):
            livello = len(spoglia) - len(spoglia.lstrip("#"))
            contenuto = spoglia.lstrip("#").strip()
            if livello == 1:
                p = doc.add_paragraph(style="Title")
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            else:
                p = doc.add_paragraph(style=f"Heading {min(livello - 1, 4)}")
            _aggiungi_testo(p, contenuto)
            i += 1
            continue

        # --- citazione ----------------------------------------------------------------
        if spoglia.startswith(">"):
            p = doc.add_paragraph(style="Intense Quote")
            _aggiungi_testo(p, spoglia.lstrip("> ").strip())
            i += 1
            continue

        # --- elenchi ------------------------------------------------------------------
        rientro = len(riga) - len(riga.lstrip(" "))
        if spoglia.startswith(("- ", "* ")):
            p = doc.add_paragraph(style="List Bullet" if rientro < 2 else "List Bullet 2")
            _aggiungi_testo(p, spoglia[2:])
            i += 1
            continue
        if re.match(r"^\d+[.)] ", spoglia):
            p = doc.add_paragraph(style="List Number" if rientro < 2 else "List Number 2")
            _aggiungi_testo(p, re.sub(r"^\d+[.)] ", "", spoglia))
            i += 1
            continue

        # --- paragrafo ----------------------------------------------------------------
        blocco = [spoglia]
        i += 1
        while (
            i < len(righe)
            and righe[i].strip()
            and not righe[i].strip().startswith(("#", "-", "*", ">", "|", "```"))
            # Un'immagine interrompe il paragrafo anche senza riga vuota davanti.
            and not _IMMAGINE.match(righe[i].strip())
        ):
            blocco.append(righe[i].strip())
            i += 1
        p = doc.add_paragraph()
        _aggiungi_testo(p, " ".join(blocco))

    destinazione.parent.mkdir(parents=True, exist_ok=True)
    doc.save(destinazione)
    return destinazione
