"""
Test della conversione Markdown -> Word dei report settimanali (`mgp.report`).

Il documento che esce da qui va al relatore, quindi gli errori di formattazione non
sono cosmetici: una barra rovesciata rimasta a vista o un corsivo aperto per sbaglio
si notano subito e tolgono credibilita' al resto.

I casi sono tutti verificabili a mano: si scrive un frammento di Markdown e si
controlla che cosa finisce nei "run" di Word.
"""

from __future__ import annotations

import sys
from pathlib import Path

from docx import Document

sys.path.insert(0, str(Path(__file__).resolve().parents[1] / "src"))

from mgp.report import _aggiungi_testo, markdown_to_docx  # noqa: E402


def _run(testo: str) -> list[tuple[str, bool, bool]]:
    """Converte un frammento inline e restituisce (testo, grassetto, corsivo) per run."""
    p = Document().add_paragraph()
    _aggiungi_testo(p, testo)
    return [(r.text, bool(r.bold), bool(r.italic)) for r in p.runs]


# --- escape Markdown ----------------------------------------------------------------


def test_escape_asterisco_perde_la_barra():
    """`K\\*` deve arrivare in Word come `K*`, non come `K\\*`.

    E' il caso che ricorre nei report, dove la soglia si scrive K\\* per impedire che
    l'asterisco apra un corsivo.
    """
    assert _run(r"La soglia K\* vale 50 MW") == [("La soglia K* vale 50 MW", False, False)]


def test_due_escape_sulla_stessa_riga_non_diventano_corsivo():
    """Il caso che la sola rimozione della barra non risolverebbe.

    Senza protezione, il testo fra i due asterischi verrebbe letto come un corsivo e
    si perderebbero entrambi i delimitatori.
    """
    assert _run(r"K\* scende mentre K\* sale") == [
        ("K* scende mentre K* sale", False, False)
    ]


def test_escape_dentro_il_codice_resta_letterale():
    """In Markdown gli escape non agiscono dentro il codice inline.

    Chi scrive `\\*` fra apici inversi vuole vedere proprio quei due caratteri.
    """
    (testo, _, _), = _run(r"il carattere `\*` in regex")[1:2]
    assert testo == r"\*"


def test_escape_dentro_il_grassetto():
    assert _run(r"**K\* definitiva**") == [("K* definitiva", True, False)]


# --- formattazione inline gia' supportata, per non farla regredire -------------------


def test_grassetto_e_corsivo_restano_distinti():
    assert _run("testo **grasso** e *corsivo* qui") == [
        ("testo ", False, False),
        ("grasso", True, False),
        (" e ", False, False),
        ("corsivo", False, True),
        (" qui", False, False),
    ]


def test_codice_inline_non_viene_interpretato():
    """Un asterisco dentro il codice non deve aprire un corsivo."""
    assert [t for t, _, _ in _run("vedi `a * b` nel testo")] == ["vedi ", "a * b", " nel testo"]


# --- immagini -------------------------------------------------------------------------


def _png(percorso: Path, larghezza: int, altezza: int) -> Path:
    """Scrive un PNG reale delle dimensioni volute (serve un file che python-docx sappia leggere)."""
    from PIL import Image

    Image.new("RGB", (larghezza, altezza), "white").save(percorso)
    return percorso


def _con_immagine(tmp_path: Path, riga: str) -> Document:
    sorgente = tmp_path / "2026-01-01_report.md"
    sorgente.write_text(f"# Prova\n\n{riga}\n", encoding="utf-8")
    return Document(markdown_to_docx(sorgente, tmp_path / "prova.docx"))


def test_immagine_larga_viene_ridotta_alla_colonna(tmp_path):
    """Una figura piu' larga del testo si scala, mantenendo le proporzioni."""
    _png(tmp_path / "larga.png", 3000, 1500)
    doc = _con_immagine(tmp_path, "![](larga.png)")

    (figura,) = doc.inline_shapes
    sezione = doc.sections[0]
    utile = sezione.page_width - sezione.left_margin - sezione.right_margin
    assert figura.width == utile
    # 2:1 in partenza, 2:1 all'arrivo: la scalatura non ha deformato.
    assert abs(figura.width / figura.height - 2.0) < 0.01


def test_immagine_stretta_non_viene_ingrandita(tmp_path):
    """Scalare in su sfoca: una figura piccola resta della sua dimensione naturale."""
    _png(tmp_path / "stretta.png", 96, 96)
    doc = _con_immagine(tmp_path, "![](stretta.png)")

    (figura,) = doc.inline_shapes
    sezione = doc.sections[0]
    assert figura.width < sezione.page_width - sezione.left_margin - sezione.right_margin
    assert abs(figura.width / figura.height - 1.0) < 0.01


def test_didascalia_sotto_la_figura(tmp_path):
    _png(tmp_path / "f.png", 400, 300)
    doc = _con_immagine(tmp_path, "![Errore per ora del giorno](f.png)")

    assert len(doc.inline_shapes) == 1
    testi = [p.text for p in doc.paragraphs if p.text.strip()]
    assert "Errore per ora del giorno" in testi


def test_didascalia_scioglie_gli_escape(tmp_path):
    """La didascalia e' testo come un altro: `K\\*` deve arrivare come `K*`."""
    _png(tmp_path / "f.png", 400, 300)
    doc = _con_immagine(tmp_path, r"![Le soglie K\* a confronto](f.png)")

    assert any("Le soglie K* a confronto" == p.text for p in doc.paragraphs)


def test_figura_mancante_lascia_un_segnaposto_visibile(tmp_path):
    """Non deve interrompere la conversione, ma non deve nemmeno passare inosservata."""
    doc = _con_immagine(tmp_path, "![](inesistente.png)")

    assert not doc.inline_shapes
    assert any("FIGURA MANCANTE: inesistente.png" in p.text for p in doc.paragraphs)


def test_riferimento_pdf_ricade_sul_png(tmp_path):
    """Word non incorpora PDF; gli script del progetto salvano sempre anche il PNG."""
    _png(tmp_path / "figura.png", 400, 300)
    doc = _con_immagine(tmp_path, "![](figura.pdf)")

    assert len(doc.inline_shapes) == 1


def test_immagine_interrompe_il_paragrafo_precedente(tmp_path):
    """Senza riga vuota davanti, `![...]()` non deve finire dentro il testo corrente."""
    _png(tmp_path / "f.png", 400, 300)
    sorgente = tmp_path / "2026-01-01_report.md"
    sorgente.write_text("# Prova\n\nTesto che precede.\n![](f.png)\n", encoding="utf-8")
    doc = Document(markdown_to_docx(sorgente, tmp_path / "prova.docx"))

    assert len(doc.inline_shapes) == 1
    assert any(p.text == "Testo che precede." for p in doc.paragraphs)


def test_figura_cercata_anche_in_output_figure(tmp_path):
    """Il nome nudo basta: evita i `../..` che renderebbero illeggibile il Markdown."""
    from mgp import config

    nome = next(config.FIGURE_DIR.glob("*.png")).name
    doc = _con_immagine(tmp_path, f"![]({nome})")
    assert len(doc.inline_shapes) == 1


# --- documento completo -------------------------------------------------------------


def test_conversione_di_un_report_minimo(tmp_path):
    """Titoli, tabella con escape nell'intestazione e corpo del testo."""
    sorgente = tmp_path / "2026-01-01_report.md"
    sorgente.write_text(
        "# Report di prova\n"
        "\n"
        "## Risultati\n"
        "\n"
        "| Anno | K\\* |\n"
        "|---|---|\n"
        "| 2024 | **50,1** |\n"
        "\n"
        "La soglia K\\* e' scesa.\n",
        encoding="utf-8",
    )
    destinazione = markdown_to_docx(sorgente, tmp_path / "prova.docx")

    doc = Document(destinazione)
    titoli = [p.text for p in doc.paragraphs if p.style.name.startswith(("Title", "Heading"))]
    assert titoli == ["Report di prova", "Risultati"]

    tabella = doc.tables[0]
    # L'intestazione passa da un percorso diverso dal corpo: va controllata a parte.
    assert [c.text for c in tabella.rows[0].cells] == ["Anno", "K*"]
    assert [c.text for c in tabella.rows[1].cells] == ["2024", "50,1"]

    assert any("La soglia K* e' scesa." == p.text for p in doc.paragraphs)


def test_nessuna_barra_rovesciata_residua_nei_report_veri(tmp_path):
    """Controllo di regressione sui report effettivamente scritti.

    Se un report contiene un escape che il convertitore non scioglie, il difetto si
    vede solo aprendo il `.docx`: questo test lo fa fallire prima.
    """
    from mgp import config

    sorgenti = sorted(
        p for p in (config.DOCS_DIR / "report").glob("*_report.md") if not p.name.startswith("_")
    )
    assert sorgenti, "nessun report da controllare"

    for sorgente in sorgenti:
        doc = Document(markdown_to_docx(sorgente, tmp_path / f"{sorgente.stem}.docx"))
        fuori_codice = [
            r.text
            for p in doc.paragraphs
            for r in p.runs
            if r.font.name != "Consolas" and "\\" in r.text
        ]
        assert not fuori_codice, f"{sorgente.name}: barre rovesciate residue {fuori_codice[:3]}"
